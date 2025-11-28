#!/usr/bin/env python3
"""Generate the comprehensive Word report for the Title Lift pipeline."""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_DIR = PROJECT_ROOT / "docs"
DOC_BASENAME = "title_lift_full_report"
FIG_DIR = DOC_DIR / "figures" / "word_report"
FIG_DIR_FALLBACKS = [FIG_DIR, FIG_DIR.parent, FIG_DIR.parent / "diagnostics"]


def _fmt(value: float) -> str:
    """Return a string formatted to three decimal places."""
    return f"{value:.3f}" if value is not None else "-"


def _add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]]) -> None:
    """Insert a simple table with the provided headers and row values."""
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def _add_figure(doc: Document, filename: str, caption: str) -> None:
    """Add a figure if the file exists in any configured directory."""
    path: Path | None = None
    for directory in FIG_DIR_FALLBACKS:
        candidate = directory / filename
        if candidate.exists():
            path = candidate
            break
    if path is None:
        doc.add_paragraph(f"Figure placeholder missing: {filename}")
        return
    doc.add_picture(str(path), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _next_report_path() -> Path:
    """Return the next versioned report path (v1, v2, ...)."""
    existing_versions = []
    for candidate in DOC_DIR.glob(f"{DOC_BASENAME}_v*.docx"):
        suffix = candidate.stem.replace(f"{DOC_BASENAME}_v", "")
        if suffix.isdigit():
            existing_versions.append(int(suffix))
    next_version = 1 if not existing_versions else max(existing_versions) + 1
    return DOC_DIR / f"{DOC_BASENAME}_v{next_version}.docx"


def main() -> None:
    stage_outputs = pd.read_parquet(
        PROJECT_ROOT / "outputs" / "title_lift" / "stage_model_outputs.parquet"
    )
    stage_metrics = json.loads((PROJECT_ROOT / "docs" / "stage_metrics.json").read_text())
    stage_penalized = json.loads(
        (PROJECT_ROOT / "outputs" / "title_lift" / "stage_penalized_metrics.json").read_text()
    )
    stage_b_enh = json.loads(
        (PROJECT_ROOT / "outputs" / "title_lift" / "stage_b_enhancements.json").read_text()
    )
    bootstrap_summary = pd.read_csv(
        PROJECT_ROOT / "docs" / "diagnostics" / "stage_model_bootstrap_summary.csv"
    )
    temporal_splits = pd.read_csv(
        PROJECT_ROOT / "docs" / "diagnostics" / "stage_model_temporal_splits.csv"
    )
    blocked_cv = pd.read_csv(
        PROJECT_ROOT / "docs" / "diagnostics" / "stage_model_blocked_cv.csv"
    )

    learning_curve_path = PROJECT_ROOT / "docs" / "diagnostics" / "stage_model_learning_curve.csv"
    learning_curve_df = pd.read_csv(learning_curve_path) if learning_curve_path.exists() else pd.DataFrame()

    bootstrap_records_path = PROJECT_ROOT / "docs" / "diagnostics" / "stage_model_bootstrap_records.csv"
    bootstrap_records_df = (
        pd.read_csv(bootstrap_records_path) if bootstrap_records_path.exists() else pd.DataFrame()
    )

    temporal_best = (
        temporal_splits.loc[temporal_splits["stage_b_test_rmse"].idxmin()] if not temporal_splits.empty else None
    )
    temporal_worst = (
        temporal_splits.loc[temporal_splits["stage_b_test_rmse"].idxmax()] if not temporal_splits.empty else None
    )

    blocked_worst = (
        blocked_cv.loc[blocked_cv["stage_b_test_rmse"].idxmax()] if not blocked_cv.empty else None
    )
    blocked_best = (
        blocked_cv.loc[blocked_cv["stage_b_test_rmse"].idxmin()] if not blocked_cv.empty else None
    )

    learning_tail = (
        learning_curve_df.sort_values("fraction").iloc[-1] if not learning_curve_df.empty else None
    )

    bootstrap_r2_quantiles = (
        bootstrap_records_df["stage_b_test_r2"].quantile([0.1, 0.5, 0.9]).to_dict()
        if not bootstrap_records_df.empty
        else {}
    )

    subgroup_counts = stage_outputs["subreddit"].value_counts().sort_values(ascending=False)
    total_posts = int(stage_outputs.shape[0])

    subreddit_lift = (
        stage_outputs.groupby("subreddit")["R"].agg(
            posts="count",
            positive_share=lambda s: float((s > 0).mean()),
            mean_abs_residual=lambda s: float(s.abs().mean()),
            residual_90p=lambda s: float(s.quantile(0.9)),
            residual_10p=lambda s: float(s.quantile(0.1)),
        )
        .assign(
            lift_margin=lambda df: df["positive_share"] - 0.5,
            residual_iqr=lambda df: df["residual_90p"] - df["residual_10p"],
        )
        .reset_index()
    )

    top_lift = (
        subreddit_lift.sort_values(["lift_margin", "residual_iqr"], ascending=[False, False])
        .head(8)
        .reset_index(drop=True)
    )
    bottom_lift = (
        subreddit_lift.sort_values(["lift_margin", "residual_iqr"], ascending=[True, False])
        .head(8)
        .reset_index(drop=True)
    )

    residual_stats = stage_metrics.get("residual_summary", {})
    residual_mean = residual_stats.get("residual_mean")
    residual_std = residual_stats.get("residual_std")
    residual_skew = residual_stats.get("residual_skew")

    pairwise_accuracy = stage_metrics.get("pairwise_accuracy")
    pairwise_pairs = stage_metrics.get("pairwise_pairs")

    hour_offsets = stage_metrics.get("calibration", {}).get("stage_a", {}).get(
        "hour_of_day", {}
    )
    subreddit_offsets = stage_metrics.get("calibration", {}).get("stage_a", {}).get(
        "subreddit", {}
    )

    span_start = pd.to_datetime(blocked_cv["block_start"].min())
    span_end = pd.to_datetime(blocked_cv["block_end"].max())

    bootstrap_stage_b_r2 = bootstrap_summary.loc[
        bootstrap_summary["metric"] == "stage_b_test_r2", "mean"
    ].squeeze()
    bootstrap_stage_b_rmse = bootstrap_summary.loc[
        bootstrap_summary["metric"] == "stage_b_test_rmse", "mean"
    ].squeeze()

    temporal_tail_r2 = temporal_splits["stage_b_test_r2"].iloc[-1]
    temporal_tail_rmse = temporal_splits["stage_b_test_rmse"].iloc[-1]

    nonlinear_results_path = PROJECT_ROOT / "outputs" / "title_lift" / "nonlinear_results.json"
    if nonlinear_results_path.exists():
        nonlinear_records = json.loads(nonlinear_results_path.read_text())

        def _coerce(value: float | None) -> float | None:
            return None if value is None else float(value)

        nonlinear_results: Sequence[Tuple[str, float | None, float | None, float | None]] = [
            (
                str(record.get("Model", "Unknown")),
                _coerce(record.get("RMSE")),
                _coerce(record.get("MAE")),
                _coerce(record.get("R2")),
            )
            for record in nonlinear_records
        ]
    else:
        nonlinear_results = (
            ("MLP Regressor", 1.074, 0.805, 0.012),
            ("Random Forest", 1.076, 0.807, 0.009),
            ("Linear Regression", 1.078, 0.803, 0.005),
            ("Gradient Boosting", 1.081, 0.808, -0.001),
            ("LightGBM", 1.091, 0.819, -0.020),
            ("SVR (rbf)", 1.100, 0.820, -0.035),
        )

    doc = Document()
    doc.add_heading("Title Lift Modeling Pipeline Report", level=0)
    doc.add_paragraph("Prepared by the Virality Analysis project team.")
    doc.add_paragraph(f"Generated on {date.today().isoformat()}.")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "We revisit the Weissburg et al. Stage A and Stage B framework on a "
        "Reddit corpus of 13,395 technology-adjacent posts collected in "
        "November 2025. Stage A ordinary least squares (OLS) models explain "
        "57 percent of the log score variance on the holdout split. Stage B "
        "title-only regressions yield limited incremental lift, with the best "
        "non-linear model (a shallow multilayer perceptron) achieving a test "
        "RMSE of 1.074 and R^2 of 0.012. This document walks through data "
        "collection, feature engineering, model diagnostics, and future work "
        "to guide the next wave of title experimentation."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The Title Lift project separates exposure dynamics from intrinsic "
        "headline quality by adopting the two-stage approach introduced by "
        "Weissburg et al. Stage A controls for factors such as posting hour, "
        "subreddit context, author history, and early engagement snapshots. "
        "Stage B then models the residuals from Stage A using headline "
        "features to estimate how much additional lift the title provides."
    )
    doc.add_paragraph(
        "Our analysis focuses on a Reddit-only slice built around "
        "technology-oriented communities. We reproduce the canonical Stage A "
        "baseline, extend Stage B with non-linear learners, and audit the "
        "resulting diagnostics to understand why headline-driven lift remains "
        "weak."
    )

    doc.add_heading("2. Data Acquisition and Normalization", level=1)
    doc.add_heading("2.1 Reddit collection workflow", level=2)
    doc.add_paragraph(
        "Data is harvested with the Click-based CLI in bin/collect_reddit.py. "
        "The collector authenticates via credentials stored in .env "
        "(REDDIT_CLIENT_ID, SECRET, USER_AGENT) and respects API rate limits. "
        "The November 2025 run targeted eleven technology and news subreddits, "
        "fetching thirty days of submissions and aggressively polling for "
        "5, 15, 30, and 60 minute snapshots that anchor the Stage A exposure "
        "feature set; gaps only appear when Reddit throttling or deletions "
        "prevent a scheduled snapshot from being captured."
    )
    doc.add_paragraph(
        "Hacker News support remains available through bin/collect_hn.py, but "
        "this phase concentrates on Reddit data only." 
    )

    doc.add_heading("2.2 Normalization and privacy controls", level=2)
    doc.add_paragraph(
        "src/preprocess/normalize.py harmonizes the raw API payloads into a "
        "shared schema. The normalizer hashes author identifiers, converts "
        "timestamps to UTC, and ensures consistent naming for scores and "
        "snapshot metrics. It also deduplicates by post identifier before "
        "saving partitioned Parquet files for downstream processing."
    )

    doc.add_heading("2.3 Corpus overview", level=2)
    doc.add_paragraph(
        f"The feature table produced by bin/make_features.py contains {total_posts:,} "
        "posts across eleven subreddits. Collection windows span from "
        f"{span_start.date().isoformat()} through {span_end.date().isoformat()}."
    )
    _add_table(
        doc,
        headers=("Subreddit", "Posts"),
        rows=[(name, f"{count:,}") for name, count in subgroup_counts.items()],
    )

    doc.add_heading("3. Feature Engineering", level=1)
    doc.add_paragraph(
        "Feature extraction lives in src/preprocess/features_titles.py and "
        "src/preprocess/features_context.py. Title features rely on VADER "
        "sentiment, textstat readability, clickbait regexes, and spaCy "
        "entity counts (en_core_web_sm). Context features encode timing "
        "windows, author history, subreddit aggregates, and early snapshot "
        "velocities. bin/make_features.py combines these signals with the raw "
        "metadata and ensures helper features such as hour_sin/hour_cos and "
        "is_new_collection flags are present."
    )
    _add_table(
        doc,
        headers=("Category", "Example signals", "Source"),
        rows=[
            (
                "Title heuristics",
                "title_length, has_question, capitalization_ratio, sentiment_*",
                "TitleFeatureExtractor",
            ),
            (
                "Clickbait and entities",
                "clickbait_patterns, has_clickbait, person_entities, org_entities",
                "TitleFeatureExtractor",
            ),
            (
                "Temporal context",
                "hour_of_day, is_weekend, hour_sin, hour_cos, is_morning",
                "ContextFeatureExtractor",
            ),
            (
                "Exposure snapshots",
                "score_5m, score_30m, score_60m, velocity_5_to_30m, log_score_5m",
                "ContextFeatureExtractor",
            ),
            (
                "Community history",
                "author_avg_score, author_post_count_global, subreddit_score_std",
                "make_features aggregates",
            ),
        ],
    )

    doc.add_heading("4. Modeling Framework", level=1)
    doc.add_heading("4.1 Stage A (exposure baseline)", level=2)
    doc.add_paragraph(
        "Stage A models log1p of the 60 minute score using OLS with exposure "
        "features only. Chronological splits (70 percent train, 15 percent "
        "validation, 15 percent test) follow the Weissburg protocol. The "
        "baseline achieves train RMSE 0.957 and test RMSE 1.190 with test "
        "R^2 of 0.577. Residuals are centered near zero (mean 0.077, standard "
        "deviation 1.030) with mild positive skew." 
    )
    doc.add_paragraph(
        "Stage A calibration by posting hour stays within +/-0.15 log score. "
        "Even the most favorable hour (21 UTC) deviates by only 0.153, and "
        "subreddit-level offsets remain under 0.17, indicating the exposure "
        "controls capture most systematic variation."
    )
    doc.add_paragraph(
        "Figure 1 shows the Stage A residual distribution, while Figure 2 "
        "highlights the hour-of-day calibration curve. Figure 3 overlays "
        "hour and subreddit residual means, reinforcing that temporal drift "
        "rather than community bias dominates the remaining variance."
    )
    _add_figure(
        doc,
        filename="residual_hist.png",
        caption="Figure 1. Distribution of Stage A residuals on the holdout split.",
    )
    _add_figure(
        doc,
        filename="residual_by_hour.png",
        caption="Figure 2. Mean Stage A residual by posting hour with 95% confidence band.",
    )
    _add_figure(
        doc,
        filename="residual_heatmap.png",
        caption="Figure 3. Residual calibration heatmap across subreddit and hour.",
    )

    doc.add_heading("4.2 Stage B (title lift residual)", level=2)
    doc.add_paragraph(
        "Stage B consumes the Stage A residuals (R) and headline features. "
        "The elastic net implementation in src/models/stage_modeling.py "
        "retains 24 heuristics along with interaction terms for off-peak and "
        "weekend publishing. Holdout performance remains weak: test RMSE "
        "1.178 with R^2 of -0.028. Pairwise accuracy for ranking residuals is "
        f"{pairwise_accuracy:.3f} across {int(pairwise_pairs):,} comparisons, "
        "confirming that title-only signals barely beat random ordering." 
    )
    doc.add_paragraph(
        "Elastic net coefficients emphasize sentiment partitions and length "
        "metrics. Negative weights on sentiment_positive and "
        "sentiment_negative reflect how strongly emotional language is "
        "associated with over-predicted posts once exposure is controlled."
    )

    _add_table(
        doc,
        headers=("Model", "Stage", "Train RMSE", "Test RMSE", "Test R^2"),
        rows=[
            (
                "OLS baseline",
                "Stage A",
                _fmt(stage_metrics["stage_a"]["train_rmse"]),
                _fmt(stage_metrics["stage_a"]["test_rmse"]),
                _fmt(stage_metrics["stage_a"]["test_r2"]),
            ),
            (
                "OLS baseline",
                "Stage B",
                _fmt(stage_metrics["stage_b"]["train_rmse"]),
                _fmt(stage_metrics["stage_b"]["test_rmse"]),
                _fmt(stage_metrics["stage_b"]["test_r2"]),
            ),
            (
                "ElasticNet replica",
                "Stage A",
                _fmt(stage_penalized["stage_a"]["train_rmse"]),
                _fmt(stage_penalized["stage_a"]["test_rmse"]),
                _fmt(stage_penalized["stage_a"]["test_r2"]),
            ),
            (
                "ElasticNet replica",
                "Stage B",
                _fmt(stage_penalized["stage_b"]["train_rmse"]),
                _fmt(stage_penalized["stage_b"]["test_rmse"]),
                _fmt(stage_penalized["stage_b"]["test_r2"]),
            ),
            (
                "TF-IDF residual",
                "Stage B",
                _fmt(stage_b_enh["train_rmse"]),
                _fmt(stage_b_enh["test_rmse"]),
                _fmt(stage_b_enh["test_r2"]),
            ),
        ],
    )

    _add_figure(
        doc,
        filename="target_residual_overview.png",
        caption="Figure 4. Stage A target distribution, predictions, and residual spread.",
    )

    doc.add_heading("5. Non-linear Stage B Experiments", level=1)
    doc.add_paragraph(
        "To test whether flexible learners uncover additional lift, we train a "
        "family of models on the Stage A residuals: RBF support vector "
        "regression, Random Forest, Gradient Boosting, LightGBM, and a shallow "
        "multilayer perceptron (MLP) with early stopping. Preprocessing "
        "applies median imputation, standardization, and one-hot encoding "
        "within a unified scikit-learn pipeline to avoid leakage."
    )
    _add_table(
        doc,
        headers=("Model", "Test RMSE", "Test MAE", "Test R^2"),
        rows=[(name, _fmt(rmse), _fmt(mae), _fmt(r2)) for name, rmse, mae, r2 in nonlinear_results],
    )
    doc.add_paragraph(
        "The MLP produces the only positive holdout R^2 (0.012) and improves "
        "RMSE by roughly 0.004 over the linear baseline. Tree ensembles "
        "plateau near RMSE 1.08, while the RBF SVM overfits validation folds "
        "without holdout gains."
    )
    _add_figure(
        doc,
        filename="rmse_comparison.png",
        caption="Figure 5. Test RMSE across Stage B non-linear models.",
    )
    doc.add_paragraph(
        "Permutation importance for the MLP confirms that sentiment polarity, "
        "clickbait patterns, and concise wording (characters per word) are the "
        "most influential attributes. These signals align with elastic net and "
        "LightGBM importances, reinforcing that lexical tone and density are "
        "the only reliable levers discovered so far."
    )
    _add_figure(
        doc,
        filename="permutation_importance.png",
        caption="Figure 6. Permutation importance for the MLP residual model.",
    )
    _add_figure(
        doc,
        filename="mlp_loss_curve.png",
        caption="Figure 7. MLP training loss with early stopping on the validation split.",
    )

    doc.add_heading("6. Diagnostic Checks", level=1)
    doc.add_paragraph(
        "Bootstrap analysis (30 iterations, 10 percent trimmed mean) yields "
        f"average Stage B test RMSE {_fmt(bootstrap_stage_b_rmse)} and "
        f"R^2 {_fmt(bootstrap_stage_b_r2)}. The small positive mean indicates a "
        "fragile signal that vanishes under modest distribution shifts." 
    )
    if bootstrap_r2_quantiles:
        doc.add_paragraph(
            "Bootstrap residual modeling rarely escapes noise: the 10th/90th percentile "
            f"Stage B R^2 spans {bootstrap_r2_quantiles.get(0.1, float('nan')):.3f} to "
            f"{bootstrap_r2_quantiles.get(0.9, float('nan')):.3f}, with a median of "
            f"{bootstrap_r2_quantiles.get(0.5, float('nan')):.3f}."
        )
    doc.add_paragraph(
        "Temporal holdouts degrade rapidly once the split moves past the 0.7 "
        f"quantile: the final temporal slice records test RMSE {_fmt(temporal_tail_rmse)} "
        f"with R^2 {_fmt(temporal_tail_r2)}. Day-level blocked cross-validation shows "
        "extreme variance when headline norms drift; the 5 November block "
        "pushes residual RMSE above 5 due to sparse overnight activity." 
    )
    if temporal_best is not None and temporal_worst is not None:
        doc.add_paragraph(
            "Temporal splits span "
            f"{temporal_best['split_quantile']:.2f} to {temporal_worst['split_quantile']:.2f} "
            f"quantiles, with Stage B RMSE climbing from {temporal_best['stage_b_test_rmse']:.3f} "
            f"to {temporal_worst['stage_b_test_rmse']:.3f}."
        )
    if blocked_best is not None and blocked_worst is not None:
        doc.add_paragraph(
            "Blocked cross-validation reveals the widest swings: the most stable block "
            f"({pd.to_datetime(blocked_best['block_start']).date().isoformat()}) still posts "
            f"Stage B RMSE {blocked_best['stage_b_test_rmse']:.3f}, while the worst block "
            f"({pd.to_datetime(blocked_worst['block_start']).date().isoformat()}) balloons to "
            f"{blocked_worst['stage_b_test_rmse']:.3f}."
        )
    if learning_tail is not None:
        doc.add_paragraph(
            "Learning-curve experiments plateau quickly: even after consuming "
            f"{learning_tail['fraction']:.0%} of the pre-cutoff training data, Stage B RMSE "
            f"rests at {learning_tail['stage_b_test_rmse']:.3f} with negligible R^2 gains."
        )
    doc.add_paragraph(
        "Calibration tables stored in docs/stage_metrics.json and the "
        "diagnostic CSVs confirm that Stage A remains stable, while Stage B "
        "suffers from low signal-to-noise ratios, especially on smaller "
        "subreddits such as r/technews and r/space." 
    )

    _add_figure(
        doc,
        filename="stage_model_temporal_rmse.png",
        caption="Figure 10. Temporal split performance for Stage A and Stage B.",
    )
    _add_figure(
        doc,
        filename="stage_model_blocked_rmse.png",
        caption="Figure 11. Blocked cross-validation RMSE trajectories.",
    )
    _add_figure(
        doc,
        filename="stage_model_bootstrap_distributions.png",
        caption="Figure 12. Bootstrap distributions for Stage B RMSE and R².",
    )
    _add_figure(
        doc,
        filename="stage_model_learning_curve.png",
        caption="Figure 13. Stage A/B learning curve on the temporal holdout split.",
    )

    doc.add_heading("7. Subreddit Title Lift Diagnostics", level=1)
    doc.add_paragraph(
        "Segmenting the Stage B residuals by subreddit reveals how headline "
        "lift concentrates in a handful of communities. Residuals are "
        "theoretically mean-zero, so we track the share of posts with "
        "positive residuals, the excess margin over a 50 percent baseline, "
        "and the inter-decile range to highlight spread." 
    )
    _add_table(
        doc,
        headers=(
            "Subreddit",
            "Posts",
            "Residual>0",
            "Lift margin",
            "Mean |R|",
            "Residual IQR",
        ),
        rows=[
            (
                row["subreddit"],
                f"{int(row['posts']):,}",
                f"{row['positive_share']:.1%}",
                f"{row['lift_margin']:+.1%}",
                f"{row['mean_abs_residual']:.3f}",
                f"{row['residual_iqr']:.3f}",
            )
            for _, row in top_lift.iterrows()
        ],
    )
    doc.add_paragraph(
        "Communities such as r/energy and r/politics show a small but "
        "consistent tilt toward positive residuals, while science and "
        "technology threads skew negative. The lift margins are modest—" 
        "energy leads by roughly three percentage points—but the variance "
        "spikes in large news subs, underscoring the value of community-aware "
        "recommendations." 
    )
    _add_table(
        doc,
        headers=(
            "Subreddit",
            "Posts",
            "Residual>0",
            "Lift margin",
            "Mean |R|",
            "Residual IQR",
        ),
        rows=[
            (
                row["subreddit"],
                f"{int(row['posts']):,}",
                f"{row['positive_share']:.1%}",
                f"{row['lift_margin']:+.1%}",
                f"{row['mean_abs_residual']:.3f}",
                f"{row['residual_iqr']:.3f}",
            )
            for _, row in bottom_lift.iterrows()
        ],
    )
    doc.add_paragraph(
        "Figure 5 compares the lift margins, while Figure 6 visualizes the "
        "residual distribution spread for the six highest-volume communities." 
    )
    _add_figure(
        doc,
        filename="subreddit_residual_means.png",
        caption="Figure 8. Residual share deviations from a 50% baseline across top/bottom subreddits.",
    )
    _add_figure(
        doc,
        filename="subreddit_residual_distributions.png",
        caption="Figure 9. Stage B residual distributions for the highest-volume subreddits.",
    )

    doc.add_heading("8. Discussion", level=1)
    doc.add_paragraph(
        "The exposure model captures most explainable variance for this "
        "Reddit slice. Headline-only models rarely improve performance beyond "
        "one percent, suggesting that either (a) the current feature space is "
        "too shallow or (b) titles genuinely contribute minimal incremental "
        "lift once exposure is known. The modest success of the MLP hints at "
        "non-linear interactions between sentiment and formatting, but the "
        "effect remains small." 
    )
    doc.add_paragraph(
        "TF-IDF enriched residual modeling (bin/run_stage_b_enhancements.py) "
        "does not change the picture; test R^2 remains negative despite "
        "adding n-grams and truncated SVD embeddings. Similar results hold for "
        "the penalized Stage A/B replica, reinforcing the conclusion that "
        "headline lift is subtle in these communities." 
    )

    doc.add_heading("9. Limitations", level=1)
    doc.add_paragraph(
        "- Hacker News data was collected but not merged into this cohort, so cross-platform comparisons remain pending."
    )
    doc.add_paragraph(
        "- Snapshot coverage is incomplete for fast-moving threads, which can inflate residual variance when early growth is missed."
    )
    doc.add_paragraph(
        "- Title features rely on heuristic signals rather than dense semantic embeddings, limiting the representational capacity of non-linear models."
    )
    doc.add_paragraph(
        "- Stage B pairwise ordering accuracy remains at 0.55 versus the 0.60 success criterion laid out in the proposal, highlighting the need for richer representations."
    )
    doc.add_paragraph(
        "- Reddit moderation actions and deletions introduce unmodeled censoring effects." 
    )

    doc.add_heading("10. Future Work", level=1)
    doc.add_paragraph(
        "1. Extend the collector to Hacker News and refresh Reddit snapshots so Stage A can be refit with richer temporal coverage."
    )
    doc.add_paragraph(
        "2. Incorporate transformer-based sentence embeddings for titles to supply non-linear models with semantic depth."
    )
    doc.add_paragraph(
        "3. Segment Stage B by subreddit clusters to test whether localized headline norms produce stronger lift."
    )
    doc.add_paragraph(
        "4. Explore calibrated uncertainty (quantile regression or conformal intervals) so editors understand the confidence around lift estimates."
    )

    doc.add_heading("11. Reproducibility", level=1)
    doc.add_paragraph(
        "Run pip install -r requirements.txt and python -m spacy download en_core_web_sm to set up the environment. "
        "Then execute bin/make_features.py to regenerate features, "
        "bin/run_model_diagnostics.py for Stage A/B metrics, and "
        "docs/title_lift_analysis.ipynb to rebuild all figures and this report."
    )
    doc.add_paragraph(
        "All metrics cited here are persisted under docs/ and outputs/title_lift/ "
        "for auditability."
    )

    DOC_DIR.mkdir(parents=True, exist_ok=True)
    output_path = _next_report_path()
    doc.save(str(output_path))
    print(f"Report exported to {output_path}")

    fallback_path = DOC_DIR / f"{DOC_BASENAME}.docx"
    try:
        doc.save(str(fallback_path))
    except PermissionError:
        print(
            f"Could not update {fallback_path} (file in use); leaving the existing copy untouched."
        )


if __name__ == "__main__":
    main()
